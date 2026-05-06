"""DOCX handler — preserves H1/H2/H3 heading levels, captures paragraph text,
flattens table cells one per line."""
from __future__ import annotations

from pathlib import Path


def extract(input_path: Path) -> str:
    from docx import Document

    doc = Document(str(input_path))
    parts: list[str] = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        style = (para.style.name or "").lower()
        if style.startswith("heading 1"):
            parts.append(f"# {text}")
        elif style.startswith("heading 2"):
            parts.append(f"## {text}")
        elif style.startswith("heading 3"):
            parts.append(f"### {text}")
        else:
            parts.append(text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text:
                    parts.append(text)

    return "\n\n".join(parts) + "\n"
