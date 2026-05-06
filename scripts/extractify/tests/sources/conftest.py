"""Shared fixtures for source-handler tests. Generates minimal PDF/DOCX/EPUB
files at test time so we don't ship binary fixtures in the repo."""
from __future__ import annotations

from pathlib import Path

import pytest


@pytest.fixture
def sample_pdf(tmp_path: Path) -> Path:
    """A minimal 2-page PDF with known text per page."""
    from fpdf import FPDF

    pdf = FPDF()
    pdf.set_font("Helvetica", size=12)
    pdf.add_page()
    pdf.cell(0, 10, text="Page one greeting.")
    pdf.add_page()
    pdf.cell(0, 10, text="Page two answer.")
    out = tmp_path / "sample.pdf"
    pdf.output(str(out))
    return out


@pytest.fixture
def sample_docx(tmp_path: Path) -> Path:
    """A DOCX with a heading hierarchy + paragraphs + a table."""
    from docx import Document

    doc = Document()
    doc.add_heading("Top Title", level=1)
    doc.add_paragraph("Intro paragraph.")
    doc.add_heading("Section A", level=2)
    doc.add_paragraph("Body of section A.")
    doc.add_heading("Subsection", level=3)
    doc.add_paragraph("Detail.")
    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "cell-1"
    table.rows[0].cells[1].text = "cell-2"
    table.rows[1].cells[0].text = "cell-3"
    table.rows[1].cells[1].text = "cell-4"
    out = tmp_path / "sample.docx"
    doc.save(str(out))
    return out


@pytest.fixture
def sample_epub(tmp_path: Path) -> Path:
    """A 2-chapter EPUB with known text per chapter."""
    from ebooklib import epub

    book = epub.EpubBook()
    book.set_identifier("test-id")
    book.set_title("Test Book")
    book.set_language("en")

    ch1 = epub.EpubHtml(title="Ch1", file_name="ch1.xhtml", lang="en")
    ch1.content = "<html><body><h1>Chapter 1</h1><p>First chapter body.</p></body></html>"
    ch2 = epub.EpubHtml(title="Ch2", file_name="ch2.xhtml", lang="en")
    ch2.content = "<html><body><h1>Chapter 2</h1><p>Second chapter body.</p></body></html>"
    book.add_item(ch1)
    book.add_item(ch2)
    book.spine = ["nav", ch1, ch2]
    book.add_item(epub.EpubNcx())
    book.add_item(epub.EpubNav())

    out = tmp_path / "sample.epub"
    epub.write_epub(str(out), book, {})
    return out
